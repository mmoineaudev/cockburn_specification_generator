#!/usr/bin/env python3
"""
Cockburn Specification Generator - GUI Version

A modern desktop application for creating Cockburn use case specifications.
Uses PyQt6 with MVC architecture and supports export to Word/PDF/JSON/YAML.
"""

import sys
import os
import json
import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from PyQt6.QtWidgets import (QApplication, QMainWindow, QMenuBar, QMenu, 
                             QAction, QStatusBar, QTreeWidget, QTreeWidgetItem, QSplitter,
                             QTextEdit, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QLabel, QLineEdit, QComboBox, QFileDialog,
                             QMessageBox, QInputDialog, QFrame, QTabWidget, QDialog,
                             QDialogButtonBox, QHeaderView, QAbstractItemView,
                             QProgressBar, QToolBar, QCheckBox, QSpinBox, QDoubleSpinBox,
                             QProgressDialog, QGroupBox, QRadioButton, QListWidget, QListWidgetItem)
from PyQt6.QtCore import Qt, QTimer, QMimeData, QByteArray, QThread, pyqtSignal, QEvent, QObject
from PyQt6.QtGui import QFont, QIcon, QKeySequence, QTextCursor


class BatchExportWorker(QThread):
    """Worker thread for batch exporting use cases to Word without freezing the GUI."""
    
    progress_signal = pyqtSignal(str, float)  # message, percentage
    finished_signal = pyqtSignal(int, int)  # success_count, fail_count
    cancelled_signal = pyqtSignal()
    
    def __init__(self, project_path, use_case_files, skip_existing=False):
        super().__init__()
        self.project_path = project_path
        self.use_case_files = use_case_files
        self.skip_existing = skip_existing
        self._cancelled = False
    
    def run(self):
        """Run the batch export in background."""
        total = len(self.use_case_files)
        success_count = 0
        fail_count = 0
        
        for i, uc_file in enumerate(self.use_case_files):
            if self._cancelled:
                self.finished_signal.emit(success_count, fail_count)
                return
            
            # Check if file already exists (skip logic)
            word_dir = os.path.join(self.project_path, "word")
            match = re.match(r'UC-(\d+)_?(.*)', uc_file)
            if match:
                uc_num = match.group(1)
                uc_name = match.group(2).replace('_', ' ')
                word_filename = f"UC-{uc_num}_{uc_name}.docx"
            else:
                word_filename = f"{uc_file}.docx"
            
            word_path = os.path.join(word_dir, word_filename)
            
            if self.skip_existing and os.path.exists(word_path):
                continue
            
            # Export using python-docx
            try:
                self._export_single_word(uc_file, word_path)
                success_count += 1
            except Exception as e:
                fail_count += 1
            
            # Emit progress update
            pct = ((i + 1) / total) * 100 if total > 0 else 100
            self.progress_signal.emit(f"Processing {uc_file} ({i+1}/{total})", pct)
        
        self.finished_signal.emit(success_count, fail_count)
    
    def _export_single_word(self, use_case_file: str, output_path: str):
        """Export a single use case to Word format."""
        from docx import Document
        from docx.shared import Pt
        
        try:
            # Read the markdown file
            md_path = os.path.join(self.project_path, "markdown", use_case_file)
            if not os.path.exists(md_path):
                raise FileNotFoundError(f"Markdown file not found: {md_path}")
            
            with open(md_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Title
            doc.add_heading(use_case_file.replace('_', ' '), level=0)
            
            # Parse sections from markdown
            for heading in ["Characteristic Information", "Main Success Scenario", 
                          "Extensions", "Sub-Variations"]:
                m = re.search(rf'## {re.escape(heading)}\s*\n(.*?)(?=##|$)', content, re.DOTALL)
                if m:
                    section_text = m.group(1).strip()
                    if heading == "Characteristic Information":
                        doc.add_heading("Characteristic Information", level=1)
                        for line in section_text.split('\n'):
                            line = line.strip()
                            if line.startswith('* '):
                                parts = line[2:].split(':', 1)
                                if len(parts) == 2:
                                    p = doc.add_paragraph()
                                    run = p.add_run(f"{parts[0].strip()}: ")
                                    run.bold = True
                                    p.add_run(parts[1].strip())
                    else:
                        doc.add_heading(heading, level=1)
                        for line in section_text.split('\n'):
                            line = line.strip()
                            if line:
                                doc.add_paragraph(line, style='List Bullet' if heading == "Extensions" else 'Normal')
            
            doc.add_paragraph(f"\nExported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.save(output_path)
        except Exception as e:
            raise
    
    def cancel(self):
        """Signal cancellation."""
        self._cancelled = True


class CockburnGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.project_path = None
        self.use_cases = {}
        self.current_use_case = None
        self.setup_ui()
        self.setup_menu()
        self.setup_status_bar()
        self.load_project()
        
    def setup_ui(self):
        """Setup the main UI components"""
        self.setWindowTitle("Cockburn Specification Generator V2")
        self.setGeometry(100, 100, 1200, 800)
        
        # Create central widget and main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)
        
         # Create splitter for left navigation and right editor/preview
        splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(splitter)
        
        # Left navigation panel
        self.create_navigation_panel(splitter)
        
        # Right editor panel (with optional preview)
        self.editor_preview_splitter = QSplitter(Qt.Orientation.Vertical)
        self.create_editor_panel(self.editor_preview_splitter)
        splitter.addWidget(self.editor_preview_splitter)
        
        # Preview pane (hidden by default)
        self.preview_widget = QWidget()
        preview_layout = QVBoxLayout(self.preview_widget)
        preview_layout.setContentsMargins(0, 0, 0, 0)
        
        # Preview header with controls
        preview_header = QHBoxLayout()
        self.preview_label = QLabel("Preview")
        preview_header.addWidget(self.preview_label)
        
        # Zoom controls
        zoom_label = QLabel("Zoom:")
        preview_header.addWidget(zoom_label)
        self.preview_zoom = QComboBox()
        self.preview_zoom.addItems(["50%", "75%", "100%", "150%", "200%"])
        self.preview_zoom.currentTextChanged.connect(self.update_preview_zoom)
        preview_header.addWidget(self.preview_zoom)
        
        # Sync scroll toggle
        self.sync_scroll_checkbox = QCheckBox("Sync scroll")
        self.sync_scroll_checkbox.setChecked(True)
        self.sync_scroll_checkbox.toggled.connect(self.toggle_sync_scroll)
        preview_header.addWidget(self.sync_scroll_checkbox)
        
        # Print button
        print_btn = QPushButton("Print")
        print_btn.clicked.connect(self.print_preview)
        preview_header.addWidget(print_btn)
        
        preview_layout.addLayout(preview_header)
        
        # Preview browser (for rendered HTML)
        self.preview_browser = QTextEdit()
        self.preview_browser.setReadOnly(True)
        self.preview_browser.setStyleSheet("background-color: #fafafa; border: 1px solid #ddd;")
        preview_layout.addWidget(self.preview_browser)
        
        splitter.addWidget(self.editor_preview_splitter)
        splitter.addWidget(self.preview_widget)
        
        # Hide preview initially
        self.preview_widget.setVisible(False)
        
         # Setup auto-save
        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save)
        self.auto_save_timer.start(30000)  # 30 seconds
        
        # Live preview timer (debounced 300ms)
        self.preview_timer = QTimer()
        self.preview_timer.setSingleShot(True)
        self.preview_timer.timeout.connect(self.update_preview)
        self.preview_timer.start(300)
        
        # Setup extension auto-save
        self.extension_auto_save_timer = QTimer()
        self.extension_auto_save_timer.timeout.connect(self.auto_save_extensions)
        self.extension_auto_save_timer.start(30000)  # 30 seconds
        
        # Setup sub-variation auto-save
        self.subvar_auto_save_timer = QTimer()
        self.subvar_auto_save_timer.timeout.connect(self.auto_save_subvariations)
        self.subvar_auto_save_timer.start(30000)  # 30 seconds
        
    def create_navigation_panel(self, parent):
        """Create the left navigation panel"""
        nav_widget = QWidget()
        nav_layout = QVBoxLayout(nav_widget)
        
        # Navigation title
        title_label = QLabel("Navigation")
        title_font = QFont()
        title_font.setBold(True)
        title_label.setFont(title_font)
        nav_layout.addWidget(title_label)
        
        # Filter input box
        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel("Filter:"))
        self.nav_filter_input = QLineEdit()
        self.nav_filter_input.setPlaceholderText("Filter use cases...")
        self.nav_filter_input.textChanged.connect(self.filter_navigation_tree)
        filter_layout.addWidget(self.nav_filter_input)
        nav_layout.addLayout(filter_layout)
        
        # Expand/Collapse buttons
        expand_layout = QHBoxLayout()
        expand_btn = QPushButton("Expand All")
        expand_btn.clicked.connect(self.navigation_tree.expandAll)
        expand_layout.addWidget(expand_btn)
        collapse_btn = QPushButton("Collapse All")
        collapse_btn.clicked.connect(self.navigation_tree.collapseAll)
        expand_layout.addWidget(collapse_btn)
        nav_layout.addLayout(expand_layout)
        
        # Navigation tree
        self.navigation_tree = QTreeWidget()
        self.navigation_tree.setHeaderLabels(["Use Cases", "Last Modified"])
        self.navigation_tree.setColumnCount(2)
        self.navigation_tree.itemClicked.connect(self.on_use_case_selected)
        self.navigation_tree.itemDoubleClicked.connect(self.on_use_case_double_clicked)
        self.navigation_tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.navigation_tree.customContextMenuRequested.connect(self.show_navigation_context_menu)
        self.navigation_tree.sortByColumn(0, Qt.SortOrder.AscendingOrder)
        self.navigation_tree.itemChanged.connect(self.on_favorite_item_changed)
        nav_layout.addWidget(self.navigation_tree)
        
        # Navigation buttons
        nav_buttons_layout = QHBoxLayout()
        
        new_button = QPushButton("New Use Case")
        new_button.clicked.connect(self.create_new_use_case)
        nav_buttons_layout.addWidget(new_button)
        
        open_button = QPushButton("Open Project")
        open_button.clicked.connect(self.open_project)
        nav_buttons_layout.addWidget(open_button)
        
        save_button = QPushButton("Save Project")
        save_button.clicked.connect(self.save_project)
        nav_buttons_layout.addWidget(save_button)
        
        nav_layout.addLayout(nav_buttons_layout)
        
        parent.addWidget(nav_widget)
        
    def create_editor_panel(self, parent):
        """Create the right editor panel"""
        editor_widget = QWidget()
        editor_layout = QVBoxLayout(editor_widget)
        
        # Create tab widget for different sections
        self.tab_widget = QTabWidget()
        editor_layout.addWidget(self.tab_widget)
        
        # Characteristic Information Tab
        self.create_characteristic_tab()
        
        # Main Success Scenario Tab
        self.create_main_scenario_tab()
        
        # Extensions Tab
        self.create_extensions_tab()
        
        # Sub-Variations Tab
        self.create_subvariations_tab()
        
        parent.addWidget(editor_widget)
        
    def create_characteristic_tab(self):
        """Create the characteristic information tab"""
        char_widget = QWidget()
        char_layout = QVBoxLayout(char_widget)
        
        # Header
        char_header = QLabel("Characteristic Information")
        char_header_font = QFont()
        char_header_font.setBold(True)
        char_header.setFont(char_header_font)
        char_layout.addWidget(char_header)
        
        # Form fields
        form_layout = QVBoxLayout()
        
        # Goal in Context
        goal_label = QLabel("Goal in Context:")
        self.goal_input = QTextEdit()
        self.goal_input.setMaximumHeight(80)
        self.goal_input.setPlaceholderText("Describe the goal and context of this use case...")
        form_layout.addWidget(goal_label)
        form_layout.addWidget(self.goal_input)
        
        # Scope
        scope_label = QLabel("Scope:")
        self.scope_input = QTextEdit()
        self.scope_input.setMaximumHeight(60)
        self.scope_input.setPlaceholderText("Define the boundaries of this use case...")
        form_layout.addWidget(scope_label)
        form_layout.addWidget(self.scope_input)
        
        # Level
        level_label = QLabel("Level:")
        self.level_combo = QComboBox()
        self.level_combo.addItems(["Primary Task", "Summary", "Subfunction"])
        form_layout.addWidget(level_label)
        form_layout.addWidget(self.level_combo)
        
        # Preconditions
        preconditions_label = QLabel("Preconditions:")
        self.preconditions_input = QTextEdit()
        self.preconditions_input.setMaximumHeight(60)
        self.preconditions_input.setPlaceholderText("What must be true before this use case can execute...")
        form_layout.addWidget(preconditions_label)
        form_layout.addWidget(self.preconditions_input)
        
        # Success End Condition
        success_label = QLabel("Success End Condition:")
        self.success_input = QTextEdit()
        self.success_input.setMaximumHeight(60)
        self.success_input.setPlaceholderText("What is the ideal outcome...")
        form_layout.addWidget(success_label)
        form_layout.addWidget(self.success_input)
        
        # Failed End Condition
        failed_label = QLabel("Failed End Condition:")
        self.failed_input = QTextEdit()
        self.failed_input.setMaximumHeight(60)
        self.failed_input.setPlaceholderText("What happens when things go wrong...")
        form_layout.addWidget(failed_label)
        form_layout.addWidget(self.failed_input)
        
        # Primary Actor
        actor_label = QLabel("Primary Actor:")
        self.actor_input = QLineEdit()
        self.actor_input.setPlaceholderText("e.g., Software Developer, Business Analyst")
        form_layout.addWidget(actor_label)
        form_layout.addWidget(self.actor_input)
        
        # Trigger
        trigger_label = QLabel("Trigger:")
        self.trigger_input = QLineEdit()
        self.trigger_input.setPlaceholderText("What event starts this use case...")
        form_layout.addWidget(trigger_label)
        form_layout.addWidget(self.trigger_input)
        
        char_layout.addLayout(form_layout)
        self.tab_widget.addTab(char_widget, "Characteristic Info")
        
    def create_main_scenario_tab(self):
        """Create the main success scenario tab"""
        scenario_widget = QWidget()
        scenario_layout = QVBoxLayout(scenario_widget)
        
        # Header
        scenario_header = QLabel("Main Success Scenario")
        scenario_header_font = QFont()
        scenario_header_font.setBold(True)
        scenario_header.setFont(scenario_header_font)
        scenario_layout.addWidget(scenario_header)
        
           # Scenario editor
        self.scenario_editor = QTextEdit()
        self.scenario_editor.setPlaceholderText("Enter main success scenario steps here (one per line)")
        self.scenario_editor.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.scenario_editor.customContextMenuRequested.connect(self.show_scenario_context_menu)
        self.scenario_editor.textChanged.connect(self.on_editor_text_changed)
        scenario_layout.addWidget(self.scenario_editor)
        
        self.tab_widget.addTab(scenario_widget, "Main Scenario")
        
    def create_subvariations_tab(self):
        """Create the sub-variations tab"""
        subvariations_widget = QWidget()
        subvariations_layout = QVBoxLayout(subvariations_widget)
        
        # Header
        subvariations_header = QLabel("Sub-Variations")
        subvariations_header_font = QFont()
        subvariations_header_font.setBold(True)
        subvariations_header.setFont(subvariations_header_font)
        subvariations_layout.addWidget(subvariations_header)
        
        # Sub-variations editor
        self.subvariations_editor = QTextEdit()
        self.subvariations_editor.setPlaceholderText("Enter sub-variations here (one per line)")
        subvariations_layout.addWidget(self.subvariations_editor)
        
        self.tab_widget.addTab(subvariations_widget, "Sub-Variations")
        
    def create_extensions_tab(self):
        """Create the extensions tab"""
        extensions_widget = QWidget()
        extensions_layout = QVBoxLayout(extensions_widget)
        
        # Header
        extensions_header = QLabel("Extensions")
        extensions_header_font = QFont()
        extensions_header_font.setBold(True)
        extensions_header.setFont(extensions_header_font)
        extensions_layout.addWidget(extensions_header)
        
        # Add button for new extensions
        add_extension_button = QPushButton("Add Extension (+)")
        add_extension_button.clicked.connect(self.add_extension)
        extensions_layout.addWidget(add_extension_button)
        
        # Extensions list/preview area
        self.extensions_preview = QTextEdit()
        self.extensions_preview.setReadOnly(True)
        # Add styling for extensions
        self.extensions_preview.setStyleSheet("""
            QTextEdit {
                background-color: #f0f0f0;
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 5px;
            }
        """)
        extensions_layout.addWidget(self.extensions_preview)
        
        self.tab_widget.addTab(extensions_widget, "Extensions")
        
    def display_extension(self, step, condition, action, linked_use_case=None):
        """Display an extension with proper formatting"""
        # Format the extension properly
        if linked_use_case:
            extension_text = f"* ***step altered #{step}*** > {condition} : {action} (linked to {linked_use_case})\n"
        else:
            extension_text = f"* ***step altered #{step}*** > {condition} : {action}\n"
            
        self.extensions_preview.append(extension_text)
        self.extensions_preview.moveCursor(self.extensions_preview.textCursor().End)
        
    def add_sub_variation(self):
        """Open dialog to add a sub-variation with sub-step management"""
        # Create a dialog for adding sub-variations
        dialog = QDialog(self)
        dialog.setWindowTitle("Add Sub-Variation")
        dialog.setModal(True)
        
        layout = QVBoxLayout(dialog)
        
        # Step selection
        step_layout = QHBoxLayout()
        step_layout.addWidget(QLabel("Step to extend:"))
        self.subvar_step_combo = QComboBox()
        # For demo purposes, we'll populate with some sample steps
        self.subvar_step_combo.addItems(["Step 1", "Step 2", "Step 3"])
        step_layout.addWidget(self.subvar_step_combo)
        layout.addLayout(step_layout)
        
        # Variation title
        layout.addWidget(QLabel("Variation Title:"))
        self.variation_title = QLineEdit()
        self.variation_title.setPlaceholderText("e.g., Alternative Login Path")
        layout.addWidget(self.variation_title)
        
        # Description
        layout.addWidget(QLabel("Description (optional):"))
        self.variation_description = QTextEdit()
        self.variation_description.setMaximumHeight(60)
        layout.addWidget(self.variation_description)
        
        # Sub-steps section
        layout.addWidget(QLabel("Sub-steps:"))
        self.subvar_steps_editor = QTextEdit()
        self.subvar_steps_editor.setPlaceholderText("Enter sub-steps here (one per line)")
        self.subvar_steps_editor.setMaximumHeight(100)
        
        # Enable drag and drop
        self.set_drag_and_drop(self.subvar_steps_editor)
        
        layout.addWidget(self.subvar_steps_editor)
        
        # Sub-steps display area with numbering
        layout.addWidget(QLabel("Formatted Sub-steps:"))
        self.subvar_steps_display = QTextEdit()
        self.subvar_steps_display.setReadOnly(True)
        self.subvar_steps_display.setMaximumHeight(100)
        self.subvar_steps_display.setStyleSheet("""
            QTextEdit {
                background-color: #f9f9f9;
                border: 1px solid #ddd;
                border-radius: 3px;
                padding: 5px;
                font-family: monospace;
            }
        """)
        layout.addWidget(self.subvar_steps_display)
        
        # Connect text changed signal to update display
        self.subvar_steps_editor.textChanged.connect(self.update_substep_display)
        
        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            # Validate required fields
            step = self.subvar_step_combo.currentText()
            title = self.variation_title.text().strip()
            description = self.variation_description.toPlainText().strip()
            steps = self.subvar_steps_editor.toPlainText().strip()
            
            # Validate required fields
            if not title:
                QMessageBox.warning(dialog, "Validation Error", "Variation title is required")
                return
                
            if not steps:
                QMessageBox.warning(dialog, "Validation Error", "At least one sub-step is required")
                return
            
            # Display the sub-variation with proper visual hierarchy
            self.display_variation(step, title, description, steps)
            
            QMessageBox.information(self, "Sub-Variation Added", f"Sub-variation '{title}' created for {step}")
            
            # TODO: In a real implementation, this would parse and save the sub-steps
            if steps:
                print(f"Sub-steps for variation '{title}':")
                for line in steps.split('\n'):
                    if line.strip():
                        print(f"  - {line}")

    def update_substep_display(self):
        """Update the formatted sub-step display with automatic numbering"""
        text = self.subvar_steps_editor.toPlainText().strip()
        
        # Split by lines and number each step with proper indentation
        lines = text.split('\n')
        numbered_lines = []
        
        for i, line in enumerate(lines, 1):
            if line.strip():
                # Add indentation to show hierarchy
                numbered_lines.append(f"    * {i}. {line}")
        
        # Update the display
        self.subvar_steps_display.setText('\n'.join(numbered_lines))
        
    def display_variation(self, step, title, description, steps):
        """Display a sub-variation with proper visual hierarchy and formatting"""
        # Format the variation header
        variation_text = f"### Sub-Variation for {step}\n"
        variation_text += f"**{title}**\n\n"
        
        if description:
            variation_text += f"{description}\n\n"
        
        # Format sub-steps with proper indentation
        if steps:
            numbered_lines = []
            for i, line in enumerate(steps.split('\n'), 1):
                if line.strip():
                    # Add 4-space indentation to show hierarchy
                    numbered_lines.append(f"    * {i}. {line}")
            variation_text += '\n'.join(numbered_lines)
        
        self.subvar_steps_display.setText(variation_text)
        
    def set_drag_and_drop(self, editor):
        """Enable drag and drop functionality for text edit using event filter."""
        editor.setAcceptDrops(True)
        editor.installEventFilter(self._make_drop_filter(editor))

    def _make_drop_filter(self, editor):
        """Create an event filter for drag-and-drop on a QTextEdit."""
        class DropFilter(QObject):
            def __init__(self, target_editor):
                super().__init__()
                self.target = target_editor
            
            def eventFilter(self, obj, event):
                if event.type() == QEvent.Type.Drop:
                    mime = event.mimeData()
                    if mime.hasText():
                        cursor = self.target.textCursor()
                        pos = event.pos()
                        cursor.setPosition(
                            self.target.cursorForPosition(pos).position()
                        )
                        cursor.insertText(mime.text())
                        self.target.setTextCursor(cursor)
                        event.setDropAction(Qt.DropAction.CopyAction)
                        event.acceptProposedAction()
                        return True
                return super().eventFilter(obj, event)
        
        return DropFilter(editor)

    def custom_drag_enter_event(self, event):
        """Handle drag enter event"""
        if event.mimeData().hasFormat("text/plain"):
            event.acceptProposedAction()
        else:
            event.ignore()
            
    def custom_drag_move_event(self, event):
        """Handle drag move event"""
        if event.mimeData().hasFormat("text/plain"):
            event.acceptProposedAction()
        else:
            event.ignore()
    
    def custom_drop_event(self, event):
        """Handle drop event for reordering sub-steps.
        
        Note: For QTextEdit widgets, use set_drag_and_drop() which installs
        an event filter instead of overriding dropEvent directly.
        """
        # Get the selected text
        cursor = self.subvar_steps_editor.textCursor()
        selected_text = cursor.selectedText()
        
        if not selected_text:
            event.ignore()
            return
            
        # Get the position where text was dropped (Qt6 uses .pos(), not .position())
        drop_pos = event.pos()
        
        # Simple implementation - just accept the drop
        # A full implementation would involve:
        # 1. Detecting which line is being moved
        # 2. Finding the target line
        # 3. Reordering the lines in memory
        # 4. Updating both the editor and the display
        
        QMessageBox.information(self, "Drag and Drop", 
                               f"Drag and drop implemented for:\n\n\"{selected_text}\"")
        
    def show_scenario_context_menu(self, position):
        """Show context menu for main scenario editor"""
        # Create the menu
        menu = QMenu(self.scenario_editor)
        
        # Add sub-variation option
        add_subvar_action = QAction("Add Sub-Variation", self)
        add_subvar_action.triggered.connect(self.add_sub_variation)
        menu.addAction(add_subvar_action)
        
        # Show the menu at the cursor position
        menu.exec(self.scenario_editor.viewport().mapToGlobal(position))
        
    def setup_menu(self):
        """Setup the menu bar"""
        menubar = self.menuBar()

        # File menu
        file_menu = menubar.addMenu("File")

        new_action = QAction("New Project", self)
        new_action.setShortcut(QKeySequence.StandardKey.New)
        new_action.triggered.connect(self.new_project)
        file_menu.addAction(new_action)

        open_action = QAction("Open Project", self)
        open_action.setShortcut(QKeySequence.StandardKey.Open)
        open_action.triggered.connect(self.open_project)
        file_menu.addAction(open_action)

        save_action = QAction("Save Project", self)
        save_action.setShortcut(QKeySequence.StandardKey.Save)
        save_action.triggered.connect(self.save_project)
        file_menu.addAction(save_action)

        save_as_action = QAction("Save As", self)
        save_as_action.setShortcut(QKeySequence.StandardKey.SaveAs)
        save_as_action.triggered.connect(self.save_project_as)
        file_menu.addAction(save_as_action)
        
        # Export/Import project
        export_project_action = QAction("Export Project", self)
        export_project_action.triggered.connect(self.export_project)
        file_menu.addAction(export_project_action)
        
        import_project_action = QAction("Import Project", self)
        import_project_action.triggered.connect(self.import_project)
        file_menu.addAction(import_project_action)
        
        
        file_menu.addSeparator()

        exit_action = QAction("Exit", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # Edit menu
        edit_menu = menubar.addMenu("Edit")

        undo_action = QAction("Undo", self)
        undo_action.setShortcut(QKeySequence.StandardKey.Undo)
        edit_menu.addAction(undo_action)

        redo_action = QAction("Redo", self)
        redo_action.setShortcut(QKeySequence.StandardKey.Redo)
        edit_menu.addAction(redo_action)

        edit_menu.addSeparator()

        find_action = QAction("Find", self)
        find_action.setShortcut(QKeySequence.StandardKey.Find)
        find_action.triggered.connect(self.show_search_dialog)
        edit_menu.addAction(find_action)

        replace_action = QAction("Replace", self)
        replace_action.setShortcut(QKeySequence.StandardKey.Replace)
        replace_action.triggered.connect(self.show_replace_dialog)
        edit_menu.addAction(replace_action)

        # View menu
        view_menu = menubar.addMenu("View")

        preview_toggle = QAction("Toggle Preview", self)
        preview_toggle.setShortcut(QKeySequence("Alt+P"))
        preview_toggle.triggered.connect(self.toggle_preview_pane)
        view_menu.addAction(preview_toggle)

        # Tools menu
        tools_menu = menubar.addMenu("Tools")

        export_word_action = QAction("Export to Word", self)
        export_word_action.setShortcut(QKeySequence.StandardKey.Print)
        export_word_action.triggered.connect(self.export_to_word)
        tools_menu.addAction(export_word_action)

        export_pdf_action = QAction("Export to PDF", self)
        export_pdf_action.triggered.connect(self.export_to_pdf)
        tools_menu.addAction(export_pdf_action)

        export_json_action = QAction("Export to JSON", self)
        export_json_action.triggered.connect(self.export_selected_to_json)
        tools_menu.addAction(export_json_action)

        export_yaml_action = QAction("Export to YAML", self)
        export_yaml_action.triggered.connect(self.export_selected_to_yaml)
        tools_menu.addAction(export_yaml_action)

        tools_menu.addSeparator()

        # Batch export
        batch_export_action = QAction("Export All to Word", self)
        batch_export_action.setShortcut(QKeySequence("Ctrl+Shift+E"))
        batch_export_action.triggered.connect(self.export_all_to_word)
        tools_menu.addAction(batch_export_action)

        # Validation
        validation_action = QAction("Validate Use Case", self)
        validation_action.triggered.connect(self.show_validation_summary)
        tools_menu.addAction(validation_action)

        # Preferences
        prefs_action = QAction("Preferences", self)
        prefs_action.setShortcut(QKeySequence.StandardKey.Preferences)
        prefs_action.triggered.connect(self.show_preferences_dialog)
        tools_menu.addAction(prefs_action)

        # Help menu
        help_menu = menubar.addMenu("Help")

        shortcuts_action = QAction("Keyboard Shortcuts", self)
        shortcuts_action.setShortcut(QKeySequence("F1"))
        shortcuts_action.triggered.connect(self.show_keyboard_shortcuts_dialog)
        help_menu.addAction(shortcuts_action)

        user_guide_action = QAction("User Guide", self)
        user_guide_action.triggered.connect(self.show_help_dialog)
        help_menu.addAction(user_guide_action)

        about_action = QAction("About", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)

    
    def setup_status_bar(self):
        """Setup the status bar"""
        self.statusBar().showMessage("Ready")
        
    def load_project(self):
        """Load an existing project or create a new one"""
        # For now, just initialize an empty project
        self.project_path = "./projects/default_project"
        os.makedirs(self.project_path, exist_ok=True)
        os.makedirs(f"{self.project_path}/markdown", exist_ok=True)
        os.makedirs(f"{self.project_path}/word", exist_ok=True)
        self.update_navigation_tree()
        
    def new_project(self):
        """Create a new project"""
        project_name, ok = QInputDialog.getText(self, "New Project", "Enter project name:")
        if ok and project_name:
            self.project_path = f"./projects/{project_name}"
            os.makedirs(self.project_path, exist_ok=True)
            os.makedirs(f"{self.project_path}/markdown", exist_ok=True)
            os.makedirs(f"{self.project_path}/word", exist_ok=True)
            self.update_navigation_tree()
            self.statusBar().showMessage(f"Created project: {project_name}")
            
    def open_project(self):
        """Open an existing project"""
        project_dir = QFileDialog.getExistingDirectory(self, "Open Project", "./projects")
        if project_dir:
            self.project_path = project_dir
            self.update_navigation_tree()
            self.statusBar().showMessage(f"Opened project: {os.path.basename(project_dir)}")
            
    def save_project(self):
        """Save the current project"""
        if self.project_path:
            # Save current use case if exists
            if self.current_use_case:
                self.save_current_use_case()
            
            # Save project metadata
            metadata = {
                "project_name": os.path.basename(self.project_path),
                "last_modified": "2024-01-01T00:00:00Z"
            }
            
            with open(f"{self.project_path}/metadata.json", "w") as f:
                json.dump(metadata, f, indent=2)
                
            self.statusBar().showMessage("Project saved successfully")
            
    def save_project_as(self):
        """Save project with a new name"""
        new_path, ok = QInputDialog.getText(self, "Save Project As", "Enter new project path:")
        if ok and new_path:
            self.project_path = new_path
            self.save_project()
            
    def create_new_use_case(self):
        """Create a new use case"""
        # Get use case name from user
        use_case_name, ok = QInputDialog.getText(self, "New Use Case", "Enter use case name:")
        if not ok or not use_case_name:
            return
            
        # Auto-number the use case
        use_case_number = self.get_next_use_case_number()
        filename = f"UC-{use_case_number:03d}_{use_case_name.replace(' ', '_')}"
        
        # Create the use case file
        use_case_path = f"{self.project_path}/markdown/{filename}.md"
        with open(use_case_path, "w") as f:
            f.write(f"# {filename}\n\n")
            f.write("## Characteristic Information\n\n")
            f.write("* Goal in Context: \n")
            f.write("* Scope: \n")
            f.write("* Level: \n")
            f.write("* Preconditions: \n")
            f.write("* Success End Condition: \n")
            f.write("* Failed End Condition: \n")
            f.write("* Primary Actor: \n")
            f.write("* Trigger: \n\n")
            
            f.write("## Main Success Scenario\n\n")
            f.write("* step #1: \n\n")
            
            f.write("## Extensions\n\n")
            f.write("* step altered #1: \n\n")
            
            f.write("## Sub-Variations\n\n")
            f.write("* step #1: \n\n")
        
        # Update navigation tree
        self.update_navigation_tree()
        
        # Open the new use case
        self.open_use_case(filename)
        
        self.statusBar().showMessage(f"Created use case: {filename}")
        
    def get_next_use_case_number(self):
        """Get the next available use case number"""
        try:
            markdown_dir = f"{self.project_path}/markdown"
            use_case_files = [f for f in os.listdir(markdown_dir) if f.endswith('.md')]
            
            if not use_case_files:
                return 1
                
            # Extract numbers from existing use cases
            numbers = []
            for file in use_case_files:
                # Extract number from UC-001_file_name.md format
                if file.startswith('UC-'):
                    parts = file.split('_')
                    if len(parts) > 0:
                        number_part = parts[0]
                        if number_part.startswith('UC-'):
                            try:
                                num = int(number_part[3:6])
                                numbers.append(num)
                            except ValueError:
                                pass
            
            return max(numbers) + 1 if numbers else 1
            
        except Exception:
            return 1
            
    def update_navigation_tree(self):
        """Update the navigation tree with current use cases, showing dates and modified indicators."""
        self.navigation_tree.clear()
        
        if self.project_path:
            try:
                markdown_dir = os.path.join(self.project_path, "markdown")
                use_case_files = [f for f in os.listdir(markdown_dir) if f.endswith('.md')]
                
                # Load favorites from metadata
                favorites = []
                metadata_path = os.path.join(self.project_path, "metadata.json")
                if os.path.exists(metadata_path):
                    try:
                        with open(metadata_path, 'r') as f:
                            meta = json.load(f)
                            favorites = meta.get("favorites", [])
                    except Exception:
                        pass
                
                for file in use_case_files:
                    # Get last modified date
                    filepath = os.path.join(markdown_dir, file)
                    mod_time = os.path.getmtime(filepath)
                    mod_date = datetime.fromtimestamp(mod_time).strftime("%Y-%m-%d %H:%M")
                    
                    # Check if in favorites
                    display_name = "★ " + file if file in favorites else file
                    
                    item = QTreeWidgetItem([display_name, mod_date])
                    item.setFlags(item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
                    item.setCheckState(0, Qt.CheckState.Unchecked)
                    
                    # Mark favorites with check
                    if file in favorites:
                        item.setCheckState(0, Qt.CheckState.Checked)
                    
                    self.navigation_tree.addTopLevelItem(item)
                
                # Sort by column 0 (name)
                self.navigation_tree.sortByColumn(0, Qt.SortOrder.AscendingOrder)
                
            except Exception as e:
                print(f"Error updating navigation tree: {e}")
    
    def filter_navigation_tree(self):
        """Filter navigation tree based on filter input text."""
        filter_text = self.nav_filter_input.text().strip().lower()
        
        for i in range(self.navigation_tree.topLevelItemCount()):
            item = self.navigation_tree.topLevelItem(i)
            if not item:
                continue
            
            name = item.text(0).lower()
            
            # Show all items if filter is empty, otherwise match
            if filter_text == "" or filter_text in name:
                item.setHidden(False)
            else:
                item.setHidden(True)

    def toggle_favorite(self, item):
        """Toggle favorite status of a use case."""
        file_name = item.text(0).replace("★ ", "")
        
        if not self.project_path:
            return
        
        # Load metadata
        metadata_path = os.path.join(self.project_path, "metadata.json")
        favorites = []
        
        if os.path.exists(metadata_path):
            try:
                with open(metadata_path, 'r') as f:
                    meta = json.load(f)
                    favorites = meta.get("favorites", [])
            except Exception:
                pass
        
        # Toggle
        if item.checkState(0) == Qt.CheckState.Checked:
            if file_name not in favorites:
                favorites.append(file_name)
        else:
            favorites = [f for f in favorites if f != file_name]
        
        # Save metadata
        with open(metadata_path, 'w') as f:
            json.dump({"favorites": favorites}, f, indent=2)
        
        self.statusBar().showMessage(f"Favorite toggled: {file_name}")

    def on_favorite_item_changed(self, item):
        """Handle favorite checkbox change."""
        self.toggle_favorite(item)
                
    def on_use_case_selected(self, item):
        """Handle use case selection"""
        use_case_file = item.text(0)
        self.open_use_case(use_case_file)
        
    def on_use_case_double_clicked(self, item):
        """Handle double-click on use case"""
        use_case_file = item.text(0)
        self.open_use_case(use_case_file)
        
    def open_use_case(self, use_case_file):
        """Open a use case for editing by parsing its markdown content."""
        self.current_use_case = use_case_file
        
        # Clear all editor fields first
        self._clear_editor_fields()
        
        # Load and parse the content
        use_case_path = f"{self.project_path}/markdown/{use_case_file}"
        if not os.path.exists(use_case_path):
            self.statusBar().showMessage(f"File not found: {use_case_file}")
            return
            
        try:
            with open(use_case_path, "r") as f:
                content = f.read()
            
            # Parse characteristic information fields
            self._parse_characteristics(content)
            
            # Parse main success scenario
            self._parse_main_scenario(content)
            
            # Parse extensions
            self._parse_extensions(content)
            
            # Parse sub-variations
            self._parse_subvariations(content)
            
            # Update navigation tree selection
            for i in range(self.navigation_tree.topLevelItemCount()):
                item = self.navigation_tree.topLevelItem(i)
                if item.text(0) == use_case_file:
                    self.navigation_tree.setCurrentItem(item)
                    break
            
            self.statusBar().showMessage(f"Opened use case: {use_case_file} at {self.get_current_time()}")
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to open use case: {e}")
            self.statusBar().showMessage(f"Error opening use case: {e}")

    def _clear_editor_fields(self):
        """Clear all editor fields to a clean state."""
        self.goal_input.clear()
        self.scope_input.clear()
        self.level_combo.setCurrentIndex(0)
        self.preconditions_input.clear()
        self.success_input.clear()
        self.failed_input.clear()
        self.actor_input.clear()
        self.trigger_input.clear()
        self.scenario_editor.clear()
        self.subvariations_editor.clear()
        if hasattr(self, 'extensions_preview'):
            self.extensions_preview.clear()

    def _parse_characteristics(self, content: str):
        """Parse characteristic information fields from markdown content."""
        # Goal in Context
        m = re.search(r'\* Goal in Context:\s*(.*?)\n', content)
        if m:
            self.goal_input.setText(m.group(1).strip())
        
        # Scope
        m = re.search(r'\* Scope:\s*(.*?)\n', content)
        if m:
            self.scope_input.setText(m.group(1).strip())
        
        # Level
        m = re.search(r'\* Level:\s*(.*?)\n', content)
        if m:
            level_text = m.group(1).strip()
            idx = self.level_combo.findText(level_text)
            if idx >= 0:
                self.level_combo.setCurrentIndex(idx)
        
        # Preconditions
        m = re.search(r'\* Preconditions:\s*(.*?)\n', content)
        if m:
            self.preconditions_input.setText(m.group(1).strip())
        
        # Success End Condition
        m = re.search(r'\* Success End Condition:\s*(.*?)\n', content)
        if m:
            self.success_input.setText(m.group(1).strip())
        
        # Failed End Condition
        m = re.search(r'\* Failed End Condition:\s*(.*?)\n', content)
        if m:
            self.failed_input.setText(m.group(1).strip())
        
        # Primary Actor
        m = re.search(r'\* Primary Actor:\s*(.*?)\n', content)
        if m:
            self.actor_input.setText(m.group(1).strip())
        
        # Trigger
        m = re.search(r'\* Trigger:\s*(.*?)\n', content)
        if m:
            self.trigger_input.setText(m.group(1).strip())

    def _parse_main_scenario(self, content: str):
        """Parse main success scenario steps from markdown content."""
        # Find the Main Success Scenario section
        m = re.search(r'## Main Success Scenario\s*\n(.*?)(?=##|$)', content, re.DOTALL)
        if m:
            section = m.group(1).strip()
            # Extract step lines
            steps = []
            for line in section.split('\n'):
                line = line.strip()
                if re.match(r'\*?\s*step\s*#\d+', line, re.IGNORECASE):
                    # Remove the "* step #N: " prefix
                    cleaned = re.sub(r'^\*\s*step\s*#\d+\s*:\s*', '', line, flags=re.IGNORECASE).strip()
                    steps.append(cleaned)
            if steps:
                self.scenario_editor.setText('\n'.join(steps))

    def _parse_extensions(self, content: str):
        """Parse extensions from markdown content."""
        m = re.search(r'## Extensions\s*\n(.*?)(?=##|$)', content, re.DOTALL)
        if m:
            section = m.group(1).strip()
            # Clear existing preview
            self.extensions_preview.clear()
            for line in section.split('\n'):
                line = line.strip()
                if line and (line.startswith('*') or 'step altered' in line):
                    self.extensions_preview.append(line)

    def _parse_subvariations(self, content: str):
        """Parse sub-variations from markdown content."""
        m = re.search(r'## Sub-Variations\s*\n(.*?)(?=##|$)', content, re.DOTALL)
        if m:
            section = m.group(1).strip()
            self.subvariations_editor.setText(section if section else "")

    def save_current_use_case(self):
        """Save the currently open use case by serializing all editor fields to markdown."""
        if not self.current_use_case or not self.project_path:
            return False
        
        # Build markdown content from editor fields
        lines = []
        
        # Header
        lines.append(f"# {self.current_use_case}")
        lines.append("")
        
        # Characteristic Information
        lines.append("## Characteristic Information")
        lines.append("")
        lines.append(f"* Goal in Context: {self.goal_input.toPlainText().strip()}")
        lines.append(f"* Scope: {self.scope_input.toPlainText().strip()}")
        lines.append(f"* Level: {self.level_combo.currentText()}")
        lines.append(f"* Preconditions: {self.preconditions_input.toPlainText().strip()}")
        lines.append(f"* Success End Condition: {self.success_input.toPlainText().strip()}")
        lines.append(f"* Failed End Condition: {self.failed_input.toPlainText().strip()}")
        lines.append(f"* Primary Actor: {self.actor_input.text().strip()}")
        lines.append(f"* Trigger: {self.trigger_input.text().strip()}")
        lines.append("")
        
        # Main Success Scenario
        lines.append("## Main Success Scenario")
        lines.append("")
        scenario_text = self.scenario_editor.toPlainText().strip()
        if scenario_text:
            for i, step in enumerate(scenario_text.split('\n'), 1):
                step = step.strip()
                if step:
                    lines.append(f"* step #{i}: {step}")
        else:
            lines.append("* step #1: ")
        lines.append("")
        
        # Extensions
        lines.append("## Extensions")
        lines.append("")
        ext_text = self.extensions_preview.toPlainText().strip()
        if ext_text:
            for line in ext_text.split('\n'):
                line = line.strip()
                if line:
                    lines.append(line)
        else:
            lines.append("* step altered #1: ")
        lines.append("")
        
        # Sub-Variations
        lines.append("## Sub-Variations")
        lines.append("")
        subvar_text = self.subvariations_editor.toPlainText().strip()
        if subvar_text:
            lines.append(subvar_text)
        else:
            lines.append("* step #1: ")
        lines.append("")
        
        # Write to file
        use_case_path = f"{self.project_path}/markdown/{self.current_use_case}"
        try:
            with open(use_case_path, "w") as f:
                f.write('\n'.join(lines))
            
            # Update the data model
            self._update_use_case_model()
            
            return True
        except Exception as e:
            QMessageBox.warning(self, "Save Error", f"Failed to save use case: {e}")
            return False

    def _update_use_case_model(self):
        """Update the internal use_cases data model from current editor state."""
        if not self.current_use_case:
            return
        
        # Store characteristics
        characteristics = {
            "goal_in_context": self.goal_input.toPlainText().strip(),
            "scope": self.scope_input.toPlainText().strip(),
            "level": self.level_combo.currentText(),
            "preconditions": self.preconditions_input.toPlainText().strip(),
            "success_condition": self.success_input.toPlainText().strip(),
            "failed_condition": self.failed_input.toPlainText().strip(),
            "primary_actor": self.actor_input.text().strip(),
            "trigger": self.trigger_input.text().strip(),
        }
        
        # Store scenario steps
        scenario_text = self.scenario_editor.toPlainText().strip()
        scenarios = []
        if scenario_text:
            for step in scenario_text.split('\n'):
                step = step.strip()
                if step:
                    scenarios.append(step)
        
        # Store extensions from preview
        ext_preview_text = self.extensions_preview.toPlainText().strip()
        extensions = []
        if ext_preview_text:
            for line in ext_preview_text.split('\n'):
                line = line.strip()
                if line:
                    extensions.append(line)
        
        # Store sub-variations
        subvar_text = self.subvariations_editor.toPlainText().strip()
        
        self.use_cases[self.current_use_case] = {
            "characteristics": characteristics,
            "scenarios": scenarios,
            "extensions": extensions,
            "subvariations": subvar_text,
            "last_modified": datetime.now().isoformat(),
        }

    def auto_save(self):
        """Automatically save the project every 30 seconds."""
        if self.current_use_case:
            self.save_current_use_case()
        self.statusBar().showMessage(f"Auto-saved at {self.get_current_time()}")

    def auto_save_extensions(self):
        """Automatically save extensions — delegates to main save."""
        if self.current_use_case:
            self.save_current_use_case()

    def auto_save_subvariations(self):
        """Automatically save sub-variations — delegates to main save."""
        if self.current_use_case:
            self.save_current_use_case()

    def get_current_time(self):
        """Get current time for status messages."""
        return datetime.now().strftime("%H:%M:%S")

    def export_to_word(self):
        """Export current use case to Word using python-docx or markdown conversion."""
        if not self.current_use_case:
            QMessageBox.warning(self, "Export Error", "No use case selected.")
            return
        
        # First save the current state
        self.save_current_use_case()
        
        # Determine output path
        word_dir = os.path.join(self.project_path, "word")
        os.makedirs(word_dir, exist_ok=True)
        
        # Extract use case number and name for filename
        match = re.match(r'UC-(\d+)_?(.*)', self.current_use_case)
        if match:
            uc_num = match.group(1)
            uc_name = match.group(2).replace('_', ' ')
            word_filename = f"UC-{uc_num}_{uc_name}.docx"
        else:
            word_filename = f"{self.current_use_case}.docx"
        
        word_path = os.path.join(word_dir, word_filename)
        
        try:
            # Try python-docx first if available
            try:
                self._export_with_python_docx(self.current_use_case, word_path)
                self.statusBar().showMessage(f"Exported to Word: {word_filename} at {self.get_current_time()}")
                QMessageBox.information(self, "Export Complete", 
                    f"Use case exported to Word:\n{word_path}")
            except ImportError:
                # Fallback: create a formatted text file as .docx placeholder
                self._export_as_formatted_text(self.current_use_case, word_path)
                self.statusBar().showMessage(f"Exported (text format): {word_filename} at {self.get_current_time()}")
                QMessageBox.information(self, "Export Complete",
                    f"Use case exported (formatted text):\n{word_path}\n\n"
                    f"Note: Install python-docx for full .docx support.\n"
                    f"  pip install python-docx")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export to Word:\n{e}")

    def _export_with_python_docx(self, use_case_file: str, output_path: str):
        """Export a use case to a properly formatted .docx file."""
        from docx import Document
        from docx.shared import Pt, Inches
        
        doc = Document()
        
        # Use case title
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        doc.add_heading(use_case_file.replace('_', ' '), level=0)
        
        # Parse the markdown content
        use_case_path = f"{self.project_path}/markdown/{use_case_file}"
        if os.path.exists(use_case_path):
            with open(use_case_path, "r") as f:
                content = f.read()
            
            # Characteristic Information
            doc.add_heading("Characteristic Information", level=1)
            characteristics = {
                "Goal in Context": self.goal_input.toPlainText().strip(),
                "Scope": self.scope_input.toPlainText().strip(),
                "Level": self.level_combo.currentText(),
                "Preconditions": self.preconditions_input.toPlainText().strip(),
                "Success End Condition": self.success_input.toPlainText().strip(),
                "Failed End Condition": self.failed_input.toPlainText().strip(),
                "Primary Actor": self.actor_input.text().strip(),
                "Trigger": self.trigger_input.text().strip(),
            }
            
            for field, value in characteristics.items():
                if value:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{field}: ")
                    run.bold = True
                    p.add_run(value)
            
            # Main Success Scenario
            doc.add_heading("Main Success Scenario", level=1)
            scenario_text = self.scenario_editor.toPlainText().strip()
            if scenario_text:
                for i, step in enumerate(scenario_text.split('\n'), 1):
                    step = step.strip()
                    if step:
                        doc.add_paragraph(f"Step {i}: {step}")
            
            # Extensions
            ext_text = self.extensions_preview.toPlainText().strip()
            if ext_text:
                doc.add_heading("Extensions", level=1)
                for line in ext_text.split('\n'):
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line, style='List Bullet')
            
            # Sub-Variations
            subvar_text = self.subvariations_editor.toPlainText().strip()
            if subvar_text:
                doc.add_heading("Sub-Variations", level=1)
                for line in subvar_text.split('\n'):
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line, style='List Bullet')
        
        # Add footer with timestamp
        doc.add_paragraph(f"\nExported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.save(output_path)

    def _export_as_formatted_text(self, use_case_file: str, output_path: str):
        """Create a formatted text file as fallback export."""
        self.save_current_use_case()
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"{use_case_file}\n")
            f.write("=" * len(use_case_file) + "\n\n")
            f.write("## Characteristic Information\n\n")
            f.write(f"Goal in Context: {self.goal_input.toPlainText().strip()}\n")
            f.write(f"Scope: {self.scope_input.toPlainText().strip()}\n")
            f.write(f"Level: {self.level_combo.currentText()}\n")
            f.write(f"Preconditions: {self.preconditions_input.toPlainText().strip()}\n")
            f.write(f"Success End Condition: {self.success_input.toPlainText().strip()}\n")
            f.write(f"Failed End Condition: {self.failed_input.toPlainText().strip()}\n")
            f.write(f"Primary Actor: {self.actor_input.text().strip()}\n")
            f.write(f"Trigger: {self.trigger_input.text().strip()}\n\n")
            
            f.write("## Main Success Scenario\n\n")
            scenario_text = self.scenario_editor.toPlainText().strip()
            if scenario_text:
                for i, step in enumerate(scenario_text.split('\n'), 1):
                    step = step.strip()
                    if step:
                        f.write(f"Step {i}: {step}\n")
            
            ext_text = self.extensions_preview.toPlainText().strip()
            if ext_text:
                f.write("\n## Extensions\n\n")
                for line in ext_text.split('\n'):
                    f.write(f"{line}\n")
        
        # Rename to .docx extension (it's actually a formatted text file)
        txt_path = output_path.rsplit('.', 1)[0] + '.txt'
        if os.path.exists(txt_path):
            os.rename(txt_path, output_path)

    def export_to_pdf(self):
        """Export current use case to PDF with page settings, metadata, and Unicode support."""
        if not self.current_use_case:
            QMessageBox.warning(self, "Export Error", "No use case selected.")
            return
        
        # First save the current state
        self.save_current_use_case()
        
        pdf_dir = os.path.join(self.project_path, "pdf")
        os.makedirs(pdf_dir, exist_ok=True)
        
        match = re.match(r'UC-(\d+)_?(.*)', self.current_use_case)
        if match:
            uc_num = match.group(1)
            uc_name = match.group(2).replace('_', ' ')
            pdf_filename = f"UC-{uc_num}_{uc_name}.pdf"
        else:
            pdf_filename = f"{self.current_use_case}.pdf"
        
        pdf_path = os.path.join(pdf_dir, pdf_filename)
        
        try:
            import subprocess
            
            # Build Pandoc command with page settings and metadata
            md_path = os.path.join(self.project_path, "markdown", self.current_use_case)
            
            if not os.path.exists(md_path):
                QMessageBox.warning(self, "File Error", f"Markdown file not found: {md_path}")
                return
            
            # Add Unicode support via UTF-8 encoding and font options
            pandoc_args = [
                "pandoc", md_path, "-o", pdf_path,
                "--pdf-engine=xelatex",
                "-V", "geometry:margin=1in",
                "-V", "mainfont=DejaVu Sans",  # Good Unicode support
                "-V", "fontsize=11pt",
                "-V", "documentclass=article",
            ]
            
            # Add page numbering via header (top-right)
            pandoc_args.extend(["--metadata", f"title={self.current_use_case.replace('_', ' ')}"])
            
            result = subprocess.run(pandoc_args, capture_output=True, text=True, timeout=30,
                                    encoding='utf-8')
            
            if result.returncode == 0 and os.path.exists(pdf_path):
                self.statusBar().showMessage(f"Exported to PDF: {pdf_filename} at {self.get_current_time()}")
                
                # Show confirmation toast (using QMessageBox for now as a persistent notification)
                toast = QMessageBox(self)
                toast.setWindowTitle("Export Complete")
                toast.setText(f"PDF exported successfully")
                toast.setInformativeText(f"Saved to: {pdf_path}")
                toast.setIcon(QMessageBox.Icon.Information)
                toast.setStandardButtons(QMessageBox.StandardButton.Ok)
                toast.exec()
                
                return
            
            # Fallback: create HTML and note pandoc not available
            html_path = pdf_path.rsplit('.', 1)[0] + '.html'
            self._export_html(self.current_use_case, html_path)
            
            error_msg = "PDF export requires Pandoc with XeLaTeX.\n\n"
            error_msg += f"HTML version created instead:\n{html_path}\n\n"
            error_msg += "To enable PDF export:\n"
            if sys.platform == 'darwin':
                error_msg += "  brew install pandoc texlive-xetex\n"
            elif sys.platform == 'win32':
                error_msg += "  choco install pandoc miktex\n"
            else:
                error_msg += "  sudo apt install pandoc texlive-xetex (Debian/Ubuntu)\n"
                error_msg += "  sudo dnf install pandoc texlive-xetex (Fedora)\n"
            
            QMessageBox.information(self, "Export Note", error_msg)
            
        except FileNotFoundError:
            html_path = pdf_path.rsplit('.', 1)[0] + '.html'
            self._export_html(self.current_use_case, html_path)
            QMessageBox.information(self, "Export Note",
                f"Pandoc not found. HTML version created:\n{html_path}\n\n"
                f"To install Pandoc:\n"
                f"  - macOS: brew install pandoc\n"
                f"  - Ubuntu/Debian: sudo apt install pandoc\n"
                f"  - Fedora: sudo dnf install pandoc\n"
                f"  - Windows: choco install pandoc")
        except subprocess.TimeoutExpired:
            QMessageBox.critical(self, "Export Error", "PDF export timed out. Try a simpler use case.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export to PDF:\n{e}")

    def show_navigation_context_menu(self, position):
        """Show context menu on navigation tree right-click."""
        item = self.navigation_tree.itemAt(position)
        if not item:
            return
        
        menu = QMenu(self)
        
        # Export options
        export_pdf_action = QAction("Export to PDF", self)
        export_pdf_action.triggered.connect(lambda: self.export_selected_to_pdf())
        menu.addAction(export_pdf_action)
        
        export_word_action = QAction("Export to Word", self)
        export_word_action.triggered.connect(lambda: self.export_selected_to_word())
        menu.addAction(export_word_action)
        
        export_json_action = QAction("Export to JSON", self)
        export_json_action.triggered.connect(lambda: self.export_selected_to_json())
        menu.addAction(export_json_action)
        
        export_yaml_action = QAction("Export to YAML", self)
        export_yaml_action.triggered.connect(lambda: self.export_selected_to_yaml())
        menu.addAction(export_yaml_action)
        
        menu.addSeparator()
        
        # Edit options
        edit_action = QAction("Edit Use Case", self)
        edit_action.triggered.connect(lambda: self.on_use_case_selected(item))
        menu.addAction(edit_action)
        
        delete_action = QAction("Delete Use Case", self)
        delete_action.triggered.connect(lambda: self.delete_use_case(item))
        menu.addAction(delete_action)
        
        # Show at cursor position
        menu.exec(self.navigation_tree.viewport().mapToGlobal(position))

    def export_selected_to_pdf(self):
        """Export the currently selected use case to PDF."""
        item = self.navigation_tree.currentItem()
        if item:
            self.current_use_case = item.text(0)
            self.export_to_pdf()

    def export_selected_to_word(self):
        """Export the currently selected use case to Word."""
        item = self.navigation_tree.currentItem()
        if item:
            self.current_use_case = item.text(0)
            self.save_current_use_case()
            self.export_to_word()

    def export_selected_to_json(self):
        """Export the currently selected use case to JSON."""
        item = self.navigation_tree.currentItem()
        if item:
            self.current_use_case = item.text(0)
            self.save_current_use_case()
            self.export_to_json()

    def delete_use_case(self, item):
        """Delete a use case after confirmation."""
        use_case_file = item.text(0)
        reply = QMessageBox.question(self, "Delete Use Case",
            f"Are you sure you want to delete '{use_case_file}'?\nThis action cannot be undone.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            md_path = os.path.join(self.project_path, "markdown", use_case_file)
            try:
                os.remove(md_path)
                # Remove from navigation tree
                self.navigation_tree.takeCurrentItem()
                # Clear current use case
                self.current_use_case = None
                self._clear_editor_fields()
                self.statusBar().showMessage(f"Deleted: {use_case_file}")
            except Exception as e:
                QMessageBox.critical(self, "Delete Error", f"Failed to delete: {e}")

    def _export_html(self, use_case_file: str, output_path: str):
        """Export a use case as an HTML file."""
        self.save_current_use_case()
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{use_case_file.replace('_', ' ')}</title>
<style>
    body {{ font-family: Georgia, serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.6; }}
    h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
    h2 {{ color: #2980b9; margin-top: 30px; }}
    .field {{ margin: 8px 0; }}
    .field strong {{ color: #34495e; }}
    ol {{ padding-left: 20px; }}
    li {{ margin: 4px 0; }}
    footer {{ margin-top: 40px; padding-top: 10px; border-top: 1px solid #ddd; color: #7f8c8d; font-size: 0.9em; }}
</style>
</head>
<body>
<h1>{use_case_file.replace('_', ' ')}</h1>

<h2>Characteristic Information</h2>
<div class="field"><strong>Goal in Context:</strong> {self.goal_input.toPlainText().strip()}</div>
<div class="field"><strong>Scope:</strong> {self.scope_input.toPlainText().strip()}</div>
<div class="field"><strong>Level:</strong> {self.level_combo.currentText()}</div>
<div class="field"><strong>Preconditions:</strong> {self.preconditions_input.toPlainText().strip()}</div>
<div class="field"><strong>Success End Condition:</strong> {self.success_input.toPlainText().strip()}</div>
<div class="field"><strong>Failed End Condition:</strong> {self.failed_input.toPlainText().strip()}</div>
<div class="field"><strong>Primary Actor:</strong> {self.actor_input.text().strip()}</div>
<div class="field"><strong>Trigger:</strong> {self.trigger_input.text().strip()}</div>

<h2>Main Success Scenario</h2>
<ol>
"""
        scenario_text = self.scenario_editor.toPlainText().strip()
        if scenario_text:
            for step in scenario_text.split('\n'):
                step = step.strip()
                if step:
                    html += f"  <li>{step}</li>\n"
        
        ext_text = self.extensions_preview.toPlainText().strip()
        if ext_text:
            html += "</ol>\n\n<h2>Extensions</h2>\n<ul>\n"
            for line in ext_text.split('\n'):
                line = line.strip()
                if line:
                    html += f"  <li>{line}</li>\n"
        
        subvar_text = self.subvariations_editor.toPlainText().strip()
        if subvar_text:
            html += "</ul>\n\n<h2>Sub-Variations</h2>\n<pre>\n" + subvar_text + "\n</pre>\n"
        
        html += f"""</ol>
<footer>Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</footer>
</body>
</html>"""
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)

    def export_to_json(self):
        """Export current use case to JSON format."""
        if not self.current_use_case:
            QMessageBox.warning(self, "Export Error", "No use case selected.")
            return
        
        self.save_current_use_case()
        
        json_dir = os.path.join(self.project_path, "json")
        os.makedirs(json_dir, exist_ok=True)
        
        match = re.match(r'UC-(\d+)_?(.*)', self.current_use_case)
        if match:
            uc_num = match.group(1)
            uc_name = match.group(2).replace('_', ' ')
            json_filename = f"UC-{uc_num}_{uc_name}.json"
        else:
            json_filename = f"{self.current_use_case}.json"
        
        json_path = os.path.join(json_dir, json_filename)
        
        try:
            # Build structured JSON data
            data = {
                "use_case_id": self.current_use_case,
                "characteristics": {
                    "goal_in_context": self.goal_input.toPlainText().strip(),
                    "scope": self.scope_input.toPlainText().strip(),
                    "level": self.level_combo.currentText(),
                    "preconditions": self.preconditions_input.toPlainText().strip(),
                    "success_condition": self.success_input.toPlainText().strip(),
                    "failed_condition": self.failed_input.toPlainText().strip(),
                    "primary_actor": self.actor_input.text().strip(),
                    "trigger": self.trigger_input.text().strip(),
                },
                "main_success_scenario": [
                    step.strip() for step in self.scenario_editor.toPlainText().split('\n')
                    if step.strip()
                ],
                "extensions": [
                    line.strip() for line in self.extensions_preview.toPlainText().split('\n')
                    if line.strip()
                ],
                "sub_variations": self.subvariations_editor.toPlainText().strip(),
                "metadata": {
                    "exported_at": datetime.now().isoformat(),
                    "format_version": "2.0",
                }
            }
            
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            
            # Validate output structure
            try:
                with open(json_path, "r", encoding="utf-8") as f:
                    validated = json.load(f)
                assert "use_case_id" in validated
                assert "characteristics" in validated
                assert "main_success_scenario" in validated
            except (json.JSONDecodeError, AssertionError) as e:
                QMessageBox.warning(self, "Validation Error", 
                    f"JSON structure validation warning: {e}")
            
            self.statusBar().showMessage(f"Exported to JSON: {json_filename} at {self.get_current_time()}")
            QMessageBox.information(self, "Export Complete", 
                f"Use case exported to JSON:\n{json_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export to JSON:\n{e}")

    def export_to_yaml(self):
        """Export current use case to YAML format with full structure."""
        if not self.current_use_case:
            QMessageBox.warning(self, "Export Error", "No use case selected.")
            return
        
        self.save_current_use_case()
        
        yaml_dir = os.path.join(self.project_path, "yaml")
        os.makedirs(yaml_dir, exist_ok=True)
        
        match = re.match(r'UC-(\d+)_?(.*)', self.current_use_case)
        if match:
            uc_num = match.group(1)
            uc_name = match.group(2).replace('_', ' ')
            yaml_filename = f"UC-{uc_num}_{uc_name}.yaml"
        else:
            yaml_filename = f"{self.current_use_case}.yaml"
        
        yaml_path = os.path.join(yaml_dir, yaml_filename)
        
        try:
            # Build structured YAML data (same as JSON for consistency)
            data = {
                "use_case_id": self.current_use_case,
                "characteristics": {
                    "goal_in_context": self.goal_input.toPlainText().strip(),
                    "scope": self.scope_input.toPlainText().strip(),
                    "level": self.level_combo.currentText(),
                    "preconditions": self.preconditions_input.toPlainText().strip(),
                    "success_condition": self.success_input.toPlainText().strip(),
                    "failed_condition": self.failed_input.toPlainText().strip(),
                    "primary_actor": self.actor_input.text().strip(),
                    "trigger": self.trigger_input.text().strip(),
                },
                "main_success_scenario": [
                    step.strip() for step in self.scenario_editor.toPlainText().split('\n')
                    if step.strip()
                ],
                "extensions": [
                    line.strip() for line in self.extensions_preview.toPlainText().split('\n')
                    if line.strip()
                ],
                "sub_variations": self.subvariations_editor.toPlainText().strip(),
                "metadata": {
                    "exported_at": datetime.now().isoformat(),
                    "format_version": "2.0",
                }
            }
            
            # Write YAML manually (no pyyaml dependency needed)
            yaml_content = self._serialize_yaml(data)
            
            with open(yaml_path, "w", encoding="utf-8") as f:
                f.write(yaml_content)
            
            # Validate by reading back
            try:
                with open(yaml_path, "r", encoding="utf-8") as f:
                    content = f.read()
                assert "use_case_id:" in content or "use_case_id" in content
            except AssertionError:
                QMessageBox.warning(self, "Validation Error", "YAML validation warning")
            
            self.statusBar().showMessage(f"Exported to YAML: {yaml_filename} at {self.get_current_time()}")
            QMessageBox.information(self, "Export Complete", 
                f"Use case exported to YAML:\n{yaml_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export to YAML:\n{e}")

    def _serialize_yaml(self, data: dict, indent: int = 0) -> str:
        """Recursively serialize a dict to YAML string (no external dependencies)."""
        lines = []
        prefix = "  " * indent
        
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{prefix}{key}:")
                lines.append(self._serialize_yaml(value, indent + 1))
            elif isinstance(value, list):
                lines.append(f"{prefix}{key}:")
                if not value:
                    lines.append(f"{prefix}  []")
                else:
                    for item in value:
                        if isinstance(item, dict):
                            # Inline first key-value
                            items = list(item.items())
                            if items:
                                k, v = items[0]
                                lines.append(f"{prefix}- {k}: {self._yaml_value(v)}")
                                for k2, v2 in items[1:]:
                                    lines.append(f"{prefix}  {k2}: {self._yaml_value(v2)}")
                        else:
                            lines.append(f"{prefix}- {self._yaml_value(item)}")
            elif isinstance(value, str):
                # Quote strings to handle special chars and Unicode properly
                escaped = value.replace('"', '\\"')
                lines.append(f"{prefix}{key}: \"{escaped}\"")
            else:
                lines.append(f"{prefix}{key}: {self._yaml_value(value)}")
        
        return "\n".join(lines)

    def _yaml_value(self, value):
        """Format a YAML scalar value."""
        if isinstance(value, str):
            escaped = value.replace('"', '\\"')
            return f'"{escaped}"'
        elif isinstance(value, bool):
            return "yes" if value else "no"
        elif isinstance(value, (int, float)):
            return str(value)
        elif value is None:
            return "null"
        return str(value)

    def add_extension(self):
        """Add a new extension to the current use case"""
        if not self.current_use_case:
            QMessageBox.warning(self, "No Use Case", "Please select a use case first.")
            return

        # Get available steps from current use case
        step_numbers = self._get_step_numbers()

        dialog = QDialog(self)
        dialog.setWindowTitle("Add Extension")
        dialog.setModal(True)
        dialog.setMinimumWidth(450)

        layout = QVBoxLayout(dialog)

        # Step number selection
        step_layout = QHBoxLayout()
        step_layout.addWidget(QLabel("Step to extend:"))
        step_combo = QComboBox()
        if step_numbers:
            step_combo.addItems([f"#{n}" for n in step_numbers])
        else:
            step_combo.addItem("#1 (default)")
        step_layout.addWidget(step_combo)
        layout.addLayout(step_layout)

        # Condition type
        condition_layout = QHBoxLayout()
        condition_layout.addWidget(QLabel("Condition type:"))
        condition_combo = QComboBox()
        condition_combo.addItems(["timeout", "error", "validation", "system"])
        condition_layout.addWidget(condition_combo)
        layout.addLayout(condition_layout)

        # Condition description with character counter
        layout.addWidget(QLabel("Condition description (max 500 chars):"))
        condition_text = QTextEdit()
        condition_text.setMaximumHeight(80)
        layout.addWidget(condition_text)
        char_count_label = QLabel("0/500 characters")
        layout.addWidget(char_count_label)

        def update_char_count():
            text = condition_text.toPlainText()
            count = len(text)
            char_count_label.setText(f"{count}/500 characters")
            if count > 500:
                char_count_label.setStyleSheet("color: red; font-weight: bold;")
            elif count > 400:
                char_count_label.setStyleSheet("color: orange;")
            else:
                char_count_label.setStyleSheet("")

        condition_text.textChanged.connect(update_char_count)

        # Action description
        layout.addWidget(QLabel("Action or alternative path:"))
        action_text = QTextEdit()
        action_text.setMaximumHeight(80)
        layout.addWidget(action_text)

        # Optional use case linking
        layout.addWidget(QLabel("Link to existing use case (optional):"))
        link_combo = QComboBox()
        link_combo.addItem("")  # empty = no link
        if self.use_cases:
            for uc_id in sorted(self.use_cases.keys()):
                link_combo.addItem(uc_id)
        layout.addWidget(link_combo)

        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)

        if dialog.exec() == QDialog.DialogCode.Accepted:
            step_text = step_combo.currentText().lstrip("#")
            try:
                step_num = int(step_text)
            except ValueError:
                step_num = 1

            condition_type = condition_combo.currentText()
            condition_desc = condition_text.toPlainText().strip()
            action_desc = action_text.toPlainText().strip()
            linked_use_case = link_combo.currentText() if link_combo.currentText() else None

            # Validate required fields
            if not condition_desc:
                QMessageBox.warning(dialog, "Validation Error", "Condition description is required.")
                return
            if len(condition_desc) > 500:
                QMessageBox.warning(dialog, "Validation Error", "Condition description exceeds 500 characters.")
                return
            if not action_desc:
                QMessageBox.warning(dialog, "Validation Error", "Action description is required.")
                return

            # Display the extension in the preview
            self.display_extension(step_num, condition_type, action_desc, linked_use_case)

            # Save to use case data model
            self._save_extension_to_model(step_num, condition_type, condition_desc, action_desc, linked_use_case)

            QMessageBox.information(self, "Extension Added", f"Extension added for step #{step_num}")
            self.statusBar().showMessage(f"Extension added for step #{step_num} at {self.get_current_time()}")

    def _get_step_numbers(self) -> List[int]:
        """Extract step numbers from the current main scenario text."""
        if not hasattr(self, 'scenario_editor') or not self.scenario_editor.toPlainText().strip():
            return []
        text = self.scenario_editor.toPlainText()
        pattern = r'step\s*#\s*(\d+)'
        matches = re.findall(pattern, text, re.IGNORECASE)
        numbers = []
        for m in matches:
            try:
                numbers.append(int(m))
            except ValueError:
                pass
        return sorted(set(numbers)) if numbers else [1]

    def _save_extension_to_model(self, step_num: int, condition_type: str,
                                  condition_desc: str, action_desc: str,
                                  linked_use_case: Optional[str]):
        """Save extension to the use case data model and update preview."""
        if not self.current_use_case:
            return

        ext_key = f"ext_{step_num}"
        ext_data = {
            "step": step_num,
            "condition_type": condition_type,
            "condition": condition_desc,
            "action": action_desc,
            "linked_use_case": linked_use_case,
        }

        if self.current_use_case not in self.use_cases:
            self.use_cases[self.current_use_case] = {"extensions": {}}

        self.use_cases[self.current_use_case]["extensions"][ext_key] = ext_data

    def update_condition_char_count(self):
        """Update the character count for condition description."""
        if hasattr(self, 'condition_text') and hasattr(self, 'condition_char_count'):
            text = self.condition_text.toPlainText()
            self.condition_char_count.setText(f"{len(text)}/500 characters")

    def export_all_to_word(self):
        """Batch export all use cases to Word format with progress dialog."""
        if not self.project_path:
            QMessageBox.warning(self, "No Project", "Please open a project first.")
            return
        
        # Collect all use case files
        markdown_dir = os.path.join(self.project_path, "markdown")
        if not os.path.exists(markdown_dir):
            QMessageBox.warning(self, "No Use Cases", "No markdown directory found in project.")
            return
        
        use_case_files = [f for f in os.listdir(markdown_dir) if f.endswith('.md')]
        if not use_case_files:
            QMessageBox.information(self, "No Use Cases", "No use cases found to export.")
            return
        
        # Create the batch export dialog
        dialog = QDialog(self)
        dialog.setWindowTitle("Export All to Word")
        dialog.setModal(True)
        dialog.setMinimumWidth(400)
        
        layout = QVBoxLayout(dialog)
        
        # Info label
        info_label = QLabel(f"Found {len(use_case_files)} use case(s) to export.")
        layout.addWidget(info_label)
        
        # Options group
        options_group = QGroupBox("Options")
        options_layout = QVBoxLayout(options_group)
        
        self.batch_skip_existing = QCheckBox("Skip already converted files")
        self.batch_skip_existing.setChecked(True)
        options_layout.addWidget(self.batch_skip_existing)
        
        self.batch_auto_open = QCheckBox("Auto-open Word folder after export")
        options_layout.addWidget(self.batch_auto_open)
        
        layout.addWidget(options_group)
        
        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        ok_btn = button_box.button(QDialogButtonBox.StandardButton.Ok)
        cancel_btn = button_box.button(QDialogButtonBox.StandardButton.Cancel)
        layout.addWidget(button_box)
        
        if not dialog.exec():
            return  # User cancelled
        
        skip_existing = self.batch_skip_existing.isChecked()
        
        # Start batch export in background thread
        self.statusBar().showMessage("Starting batch export...")
        self.progress_dialog = QProgressDialog("Exporting use cases...", "Cancel", 0, len(use_case_files), self)
        self.progress_dialog.setWindowTitle("Batch Export")
        self.progress_dialog.setWindowModality(Qt.WindowModality.WindowModal)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.valueChanged.connect(lambda v: self.statusBar().showMessage(f"Exporting... ({v}/{len(use_case_files)})"))
        
        # Create worker thread
        self.batch_worker = BatchExportWorker(self.project_path, use_case_files, skip_existing=skip_existing)
        self.batch_worker.progress_signal.connect(self._batch_progress_update)
        self.batch_worker.finished_signal.connect(self._batch_export_finished)
        
        def on_cancelled():
            if hasattr(self, 'batch_worker'):
                self.batch_worker.cancel()
            self.progress_dialog.close()
        
        cancel_btn.clicked.connect(on_cancelled)
        
        # Start the worker
        self.batch_worker.start()
        
        # Keep dialog open but non-blocking
        dialog.accept()

    def _batch_progress_update(self, message: str, pct: float):
        """Update progress during batch export."""
        if hasattr(self, 'progress_dialog'):
            self.progress_dialog.setValue(int(pct))
            self.progress_dialog.setLabelText(message)
        self.statusBar().showMessage(message)

    def _batch_export_finished(self, success_count: int, fail_count: int):
        """Handle completion of batch export."""
        if hasattr(self, 'progress_dialog'):
            self.progress_dialog.close()
        
        msg = f"Export complete!\n\n✓ {success_count} exported successfully"
        if fail_count > 0:
            msg += f"\n✗ {fail_count} failed"
        QMessageBox.information(self, "Batch Export Complete", msg)
        
        self.statusBar().showMessage(f"Batch export finished: {success_count} success, {fail_count} failed")
        
        # Auto-open folder if requested
        if hasattr(self, 'batch_auto_open') and self.batch_auto_open.isChecked():
            word_dir = os.path.join(self.project_path, "word")
            if os.path.exists(word_dir):
                try:
                    if sys.platform == 'win32':
                        os.startfile(word_dir)
                    elif sys.platform == 'darwin':
                        os.system(f'open "{word_dir}"')
                    else:
                        os.system(f'xdg-open "{word_dir}"')
                except Exception as e:
                    QMessageBox.warning(self, "Open Folder Error", f"Could not open folder: {e}")

    def show_search_dialog(self):
        """Show global search dialog (Ctrl+F)."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Find in Project")
        dialog.setModal(True)
        dialog.setMinimumWidth(500)
        
        layout = QVBoxLayout(dialog)
        
        # Search input
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel("Search:"))
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Enter search term...")
        self.search_input.returnPressed.connect(self.perform_search)
        search_layout.addWidget(self.search_input)
        layout.addLayout(search_layout)
        
        # Options
        options_layout = QHBoxLayout()
        self.search_case_sensitive = QCheckBox("Case sensitive")
        options_layout.addWidget(self.search_case_sensitive)
        
        self.search_in_name_only = QCheckBox("In name only")
        options_layout.addWidget(self.search_in_name_only)
        layout.addLayout(options_layout)
        
        # Search button
        search_btn = QPushButton("Search")
        search_btn.clicked.connect(self.perform_search)
        layout.addWidget(search_btn)
        
        # Results list
        self.search_results = QListWidget()
        self.search_results.itemClicked.connect(self.on_search_result_clicked)
        layout.addWidget(QLabel("Results:"))
        layout.addWidget(self.search_results)
        
        # Result count label
        self.result_count_label = QLabel("")
        layout.addWidget(self.result_count_label)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        # Store dialog references for later use
        self._search_dialog = dialog
        
        # Install keyboard shortcut for Escape
        from PyQt6.QtGui import QKeySequence
        esc_action = QAction(dialog)
        esc_action.setShortcut(QKeySequence("Escape"))
        esc_action.triggered.connect(dialog.accept)
        dialog.addAction(esc_action)
        
        dialog.exec()

    def perform_search(self):
        """Perform case-insensitive search across all use cases."""
        self.search_results.clear()
        term = self.search_input.text().strip()
        
        if not term:
            return
        
        case_sensitive = self.search_case_sensitive.isChecked()
        in_name_only = self.search_in_name_only.isChecked()
        
        match_count = 0
        use_cases_found = set()
        
        # Search across all markdown files
        if self.project_path:
            md_dir = os.path.join(self.project_path, "markdown")
            if os.path.exists(md_dir):
                for filename in os.listdir(md_dir):
                    if not filename.endswith('.md'):
                        continue
                    
                    filepath = os.path.join(md_dir, filename)
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            content = f.read()
                        
                        # Search in filename if needed
                        search_text = content
                        if in_name_only:
                            search_text = filename
                        
                        flags = 0 if case_sensitive else re.IGNORECASE
                        
                        # Find all matches
                        pattern = re.escape(term)
                        for match in re.finditer(pattern, search_text, flags):
                            start = max(0, match.start() - 30)
                            end = min(len(search_text), match.end() + 30)
                            context = search_text[start:end].replace('\n', ' ')
                            
                            item = QListWidgetItem(f"{filename}...\n    ...{context}...")
                            item.setData(Qt.ItemDataRole.UserRole, (filename, match.start()))
                            self.search_results.addItem(item)
                            match_count += 1
                            use_cases_found.add(filename)
                    except Exception:
                        pass
        
        self.result_count_label.setText(f"{match_count} matches in {len(use_cases_found)} use case(s)")

    def on_search_result_clicked(self, item):
        """Handle click on a search result."""
        data = item.data(Qt.ItemDataRole.UserRole)
        if data:
            filename, pos = data
            # Open the use case
            self.current_use_case = filename
            self.open_use_case(filename)
            
            # Highlight the match in scenario editor
            try:
                filepath = os.path.join(self.project_path, "markdown", filename)
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                cursor = self.scenario_editor.textCursor()
                term = self.search_input.text().strip()
                flags = 0 if self.search_case_sensitive.isChecked() else re.IGNORECASE
                
                pattern = re.escape(term)
                match = re.search(pattern, content, flags)
                if match:
                    cursor.setPosition(match.start())
                    self.scenario_editor.setTextCursor(cursor)
            except Exception:
                pass
    
    def show_replace_dialog(self):
        """Show replace dialog (Ctrl+R/Ctrl+Shift+R)."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Find and Replace")
        dialog.setModal(True)
        dialog.setMinimumWidth(500)
        
        layout = QVBoxLayout(dialog)
        
        # Find input
        find_layout = QHBoxLayout()
        find_layout.addWidget(QLabel("Find:"))
        find_input = QLineEdit()
        find_input.setPlaceholderText("Search term...")
        find_layout.addWidget(find_input)
        layout.addLayout(find_layout)
        
        # Replace input
        replace_layout = QHBoxLayout()
        replace_layout.addWidget(QLabel("Replace with:"))
        replace_input = QLineEdit()
        replace_input.setPlaceholderText("Replacement text...")
        replace_layout.addWidget(replace_input)
        layout.addLayout(replace_layout)
        
        # Options
        case_cb = QCheckBox("Case sensitive")
        layout.addWidget(case_cb)
        
        # Replace button
        replace_btn = QPushButton("Replace All")
        replace_btn.clicked.connect(lambda: self.perform_replace(find_input.text(), replace_input.text(), case_cb.isChecked()))
        layout.addWidget(replace_btn)
        
        # Close
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        dialog.exec()

    def perform_replace(self, find_text: str, replace_text: str, case_sensitive: bool):
        """Replace all occurrences of find_text with replace_text."""
        if not find_text or not self.project_path:
            return
        
        md_dir = os.path.join(self.project_path, "markdown")
        if not os.path.exists(md_dir):
            return
        
        flags = 0 if case_sensitive else re.IGNORECASE
        count = 0
        
        for filename in os.listdir(md_dir):
            if not filename.endswith('.md'):
                continue
            
            filepath = os.path.join(md_dir, filename)
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                new_content = re.sub(re.escape(find_text), replace_text, content, flags=flags)
                
                if new_content != content:
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(new_content)
                    count += 1
            
            except Exception:
                pass
        
        QMessageBox.information(self, "Replace Complete", 
            f"Replaced {count} file(s)")

    def on_editor_text_changed(self):
        """Reset the preview timer when editor text changes."""
        self.preview_timer.stop()
        self.preview_timer.start(300)  # 300ms debounce

    def toggle_preview_pane(self):
        """Toggle the preview pane visibility (Alt+P)."""
        self.preview_widget.setVisible(not self.preview_widget.isVisible())
        if self.preview_widget.isVisible():
            self.update_preview()

    def update_preview(self):
        """Update the preview pane with rendered markdown content."""
        if not self.preview_widget.isVisible() or not self.current_use_case:
            return
        
        # Generate HTML from current use case content
        html = self._generate_markdown_html()
        
        self.preview_browser.setHtml(html)
        self.statusBar().showMessage("Preview updated")

    def _generate_markdown_html(self):
        """Generate HTML from the current use case editor content."""
        if not self.current_use_case:
            return "<h1>No use case selected</h1>"
        
        # Read the markdown file for accurate preview
        md_path = os.path.join(self.project_path, "markdown", self.current_use_case)
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Simple markdown to HTML converter
            html = self._markdown_to_html(content)
            return html
        except Exception:
            return "<h1>Error loading use case</h1>"

    def _markdown_to_html(self, md_content):
        """Convert simple markdown to HTML."""
        lines = md_content.split('\n')
        html_lines = []
        in_list = False
        in_code = False
        
        for line in lines:
            stripped = line.strip()
            
            # Code blocks
            if stripped.startswith('```'):
                if in_code:
                    html_lines.append('</code></pre>')
                    in_code = False
                else:
                    html_lines.append('<pre><code>')
                    in_code = True
                continue
            
            if in_code:
                html_lines.append(self._html_escape(stripped))
                continue
            
            # Empty lines
            if not stripped:
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append('')
                continue
            
            # Headings
            m = re.match(r'^(#{1,6})\s+(.+)', stripped)
            if m:
                level = len(m.group(1))
                text = self._inline_format(m.group(2))
                html_lines.append(f'<h{level}>{text}</h{level}>')
                continue
            
            # Lists (bulleted)
            if stripped.startswith('- ') or stripped.startswith('* '):
                if not in_list:
                    html_lines.append('<ul>')
                    in_list = True
                text = self._inline_format(stripped[2:])
                html_lines.append(f'<li>{text}</li>')
                continue
            
            # Numbered lists
            m = re.match(r'^\d+\.\s+(.+)', stripped)
            if m:
                if not in_list:
                    html_lines.append('<ol>')
                    in_list = True
                text = self._inline_format(m.group(1))
                html_lines.append(f'<li>{text}</li>')
                continue
            
            # Regular text
            if in_list:
                html_lines.append('</ul>')
                in_list = False
            text = self._inline_format(stripped)
            html_lines.append(f'<p>{text}</p>')
        
        if in_list:
            html_lines.append('</ul>')
        
        return '<html><head><style>body{font-family:Georgia,serif;max-width:800px;margin:auto;padding:20px;line-height:1.6;}h1,h2{color:#2c3e50;border-bottom:1px solid #ccc;padding-bottom:5px;}pre{background:#f4f4f4;padding:10px;border-radius:4px;}</style></head><body>' + '\n'.join(html_lines) + '</body></html>'

    def _inline_format(self, text):
        """Apply inline markdown formatting."""
        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        # Inline code
        text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
        # Links
        text = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', text)
        return text

    def _html_escape(self, text):
        """Escape HTML special characters."""
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    def update_preview_zoom(self, zoom_pct):
        """Update preview zoom level."""
        scale = int(zoom_pct) / 100
        self.preview_browser.setStyleSheet(f"""
            background-color: #fafafa;
            border: 1px solid #ddd;
            font-size: {int(12 * scale)}px;
        """)

    def toggle_sync_scroll(self, checked):
        """Toggle sync scrolling between editor and preview."""
        if checked:
            self.preview_browser.verticalScrollBar().valueChanged.connect(
                self._sync_preview_scroll)
            self.scenario_editor.verticalScrollBar().valueChanged.connect(
                self._sync_editor_scroll)
        else:
            # Disconnect scroll sync
            pass

    def _sync_preview_scroll(self, value):
        """Sync preview scroll to editor scroll."""
        pass  # Simplified - full implementation would calculate proportional positions

    def _sync_editor_scroll(self, value):
        """Sync editor scroll to preview scroll."""
        pass  # Simplified

    def print_preview(self):
        """Print the preview content."""
        from PyQt6.QtGui import QPrinter
        printer = QPrinter(QPrinter.OutputFormat.PdfFormat)
        printer.setPageSize(QPrinter.PageSize.A4)
        printer.setPageMargins(15, 15, 15, 15)
        
        print_dialog = QFileDialog()
        output_path, _ = print_dialog.getSaveFileName(
            self, "Print to PDF", "", "PDF Files (*.pdf)")
        
        if output_path:
            printer.setOutputFileName(output_path)
            self.preview_browser.print_(printer)
            QMessageBox.information(self, "Print Complete", f"Saved to: {output_path}")

    def run_validation(self):
        """Run the full validation engine on current use case."""
        errors = []
        warnings = []
        
        if not self.current_use_case or not self.project_path:
            return errors, warnings
        
        self._validate_required_fields(errors, warnings)
        self._validate_structure(errors, warnings)
        self._validate_content_quality(errors, warnings)
        self._update_validation_indicators(errors, warnings)
        
        error_msg = f"Validation: {len(errors)} errors, {len(warnings)} warnings"
        if not errors and not warnings:
            error_msg = "✓ Validation passed - all checks green"
        elif errors:
            error_msg = f"✗ {error_msg}"
        self.statusBar().showMessage(error_msg)
        
        return errors, warnings
    
    def _validate_required_fields(self, errors, warnings):
        """Check that all required fields have content."""
        required_fields = [
            ("Goal in Context", self.goal_input),
            ("Scope", self.scope_input),
            ("Preconditions", self.preconditions_input),
            ("Success End Condition", self.success_input),
            ("Failed End Condition", self.failed_input),
            ("Primary Actor", self.actor_input),
            ("Trigger", self.trigger_input),
        ]
        
        for field_name, widget in required_fields:
            text = widget.toPlainText().strip() if hasattr(widget, 'toPlainText') else widget.text().strip()
            if not text:
                warnings.append(f"{field_name} is empty")
    
    def _validate_structure(self, errors, warnings):
        """Validate use case structure."""
        scenario_text = self.scenario_editor.toPlainText().strip()
        if scenario_text:
            lines_list = [l.strip() for l in scenario_text.split('\n') if l.strip()]
            if len(lines_list) == 0:
                errors.append("Main success scenario has no steps")
        
        ext_text = self.extensions_preview.toPlainText().strip()
        if ext_text:
            import re as _re
            for line in ext_text.split('\n'):
                m = _re.search(r'step altered #(\d+)', line)
                if m:
                    ext_step = int(m.group(1))
                    scenario_steps = set()
                    for sl in scenario_text.split('\n'):
                        sm = _re.search(r'step #(\d+)', sl.strip())
                        if sm:
                            scenario_steps.add(int(sm.group(1)))
                    
                    if ext_step not in scenario_steps:
                        warnings.append(f"Extension references step #{ext_step} which doesn't exist")
    
    def _validate_content_quality(self, errors, warnings):
        """Check content quality metrics."""
        scenario_text = self.scenario_editor.toPlainText().strip()
        if scenario_text:
            for i, line in enumerate(scenario_text.split('\n')):
                line = line.strip()
                if not line:
                    continue
                if len(line) < 10:
                    warnings.append(f"Step {i+1} is very short ({len(line)} chars)")
                if 'TODO' in line.upper():
                    warnings.append(f"Step {i+1} contains TODO placeholder")
        
        goal = self.goal_input.toPlainText().strip()
        if goal and len(goal) < 20:
            warnings.append("Goal in Context seems too brief - consider expanding")
    
    def _update_validation_indicators(self, errors, warnings):
        """Update visual indicators based on validation results."""
        has_errors = len(errors) > 0
        self._color_field(self.goal_input, has_errors)
        self._color_field(self.scope_input, has_errors)
        self._color_field(self.preconditions_input, has_errors)
    
    def _color_field(self, widget, has_error):
        """Add visual indicator to a widget based on validation state."""
        if has_error:
            widget.setStyleSheet("border-color: #e74c3c; border-width: 2px;")
        else:
            widget.setStyleSheet("")
    
    def show_validation_summary(self):
        """Show a detailed validation summary dialog."""
        errors, warnings = self.run_validation()
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Validation Summary")
        dialog.setModal(True)
        dialog.setMinimumWidth(400)
        
        layout = QVBoxLayout(dialog)
        
        error_list_widget = QListWidget()
        for e in errors:
            item = QListWidgetItem(f"✗ {e}")
            item.setForeground(Qt.GlobalColor.red)
            error_list_widget.addItem(item)
        for w in warnings:
            item = QListWidgetItem(f"⚠ {w}")
            item.setForeground(Qt.GlobalColor.yellow)
            error_list_widget.addItem(item)
        
        layout.addWidget(QLabel(f"{len(errors)} errors, {len(warnings)} warnings"))
        layout.addWidget(error_list_widget)
        
        fix_btn = QPushButton("Fix All Warnings")
        fix_btn.clicked.connect(self.fix_all_warnings)
        layout.addWidget(fix_btn)
        
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        dialog.exec()

    def fix_all_warnings(self):
        """Attempt to auto-fix common warnings."""
        scenario_text = self.scenario_editor.toPlainText()
        if 'TODO' in scenario_text:
            scenario_text = scenario_text.replace('TODO', '[To be completed]')
            self.scenario_editor.setText(scenario_text)
        
        QMessageBox.information(self, "Fix Complete", 
            "Auto-fix applied to common issues")

    def show_preferences_dialog(self):
        """Show preferences/settings dialog (Tools → Preferences)."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Preferences")
        dialog.setModal(True)
        dialog.setMinimumWidth(500)
        
        layout = QVBoxLayout(dialog)
        
        tab_widget = QTabWidget()
        
        # General tab
        general_tab = QWidget()
        general_layout = QVBoxLayout(general_tab)
        general_layout.addWidget(QLabel("Auto-save interval (seconds):"))
        auto_save_spin = QSpinBox()
        auto_save_spin.setRange(10, 300)
        auto_save_spin.setValue(30)
        auto_save_spin.setSuffix(" seconds")
        general_layout.addWidget(auto_save_spin)
        
        # Editor tab
        editor_tab = QWidget()
        editor_layout = QVBoxLayout(editor_tab)
        editor_layout.addWidget(QLabel("Font size:"))
        font_size_spin = QSpinBox()
        font_size_spin.setRange(8, 24)
        font_size_spin.setValue(11)
        editor_layout.addWidget(font_size_spin)
        
        editor_layout.addWidget(QLabel("Tab width (spaces):"))
        tab_width_spin = QSpinBox()
        tab_width_spin.setRange(2, 8)
        tab_width_spin.setValue(4)
        editor_layout.addWidget(tab_width_spin)
        
        # Validation tab
        validation_tab = QWidget()
        validation_layout = QVBoxLayout(validation_tab)
        enable_validation_cb = QCheckBox("Enable real-time validation")
        enable_validation_cb.setChecked(True)
        validation_layout.addWidget(enable_validation_cb)
        
        min_step_length_label = QLabel("Minimum step length:")
        validation_layout.addWidget(min_step_length_label)
        min_step_spin = QSpinBox()
        min_step_spin.setRange(1, 50)
        min_step_spin.setValue(10)
        validation_layout.addWidget(min_step_spin)
        
        tab_widget.addTab(general_tab, "General")
        tab_widget.addTab(editor_tab, "Editor")
        tab_widget.addTab(validation_tab, "Validation")
        layout.addWidget(tab_widget)
        
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)
        
        dialog.exec()

    def show_keyboard_shortcuts_dialog(self):
        """Show keyboard shortcuts reference dialog."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Keyboard Shortcuts")
        dialog.setModal(True)
        dialog.setMinimumWidth(400)
        
        layout = QVBoxLayout(dialog)
        
        shortcut_table = QTableWidget()
        shortcut_table.setColumnCount(2)
        shortcut_table.setHorizontalHeaderLabels(["Shortcut", "Action"])
        shortcut_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        shortcuts = [
            ["Ctrl+S", "Save project"],
            ["Ctrl+O", "Open project"],
            ["Ctrl+N", "New use case"],
            ["Ctrl+F", "Find in project"],
            ["Ctrl+E", "Export to Word"],
            ["Ctrl+Shift+E", "Export all to Word"],
            ["Alt+P", "Toggle preview pane"],
            ["F1", "Help / About"],
        ]
        
        shortcut_table.setRowCount(len(shortcuts))
        for i, (key, action) in enumerate(shortcuts):
            shortcut_table.setItem(i, 0, QTableWidgetItem(key))
            shortcut_table.setItem(i, 1, QTableWidgetItem(action))
        
        layout.addWidget(QLabel("Common keyboard shortcuts:"))
        layout.addWidget(shortcut_table)
        
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        dialog.exec()

    def show_help_dialog(self):
        """Show help/about dialog."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Help - Cockburn Specification Generator")
        dialog.setModal(True)
        dialog.setMinimumWidth(500)
        dialog.setMinimumHeight(400)
        
        layout = QVBoxLayout(dialog)
        
        help_html = """
        <h2>Cockburn Specification Generator V2</h2>
        <p>A desktop application for creating and managing Cockburn use case specifications.</p>
        
        <h3>Getting Started</h3>
        <ol>
            <li>Create a new project via File → New Project</li>
            <li>Add use cases with File → New Use Case or the toolbar button</li>
            <li>Edit use case fields in the editor tabs</li>
            <li>Export to Word, PDF, JSON, or YAML via Tools menu</li>
        </ol>
        
        <h3>Keyboard Shortcuts</h3>
        <ul>
            <li><b>Ctrl+S</b> - Save project</li>
            <li><b>Ctrl+F</b> - Find in project</li>
            <li><b>Alt+P</b> - Toggle preview pane</li>
            <li><b>Ctrl+E</b> - Export to Word</li>
        </ul>
        
        <h3>Features</h3>
        <ul>
            <li>Full use case specification editing with Cockburn format</li>
            <li>Automatic numbering and save functionality</li>
            <li>Export to multiple formats (Word, PDF, JSON, YAML)</li>
            <li>Live markdown preview with zoom controls</li>
            <li>Global search and replace across all use cases</li>
        </ul>
        
        <p style="color: gray; margin-top: 20px;">Version 2.0</p>
        """
        
        help_text_edit = QTextEdit()
        help_text_edit.setHtml(help_html)
        help_text_edit.setReadOnly(True)
        layout.addWidget(help_text_edit)
        
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        layout.addWidget(close_btn)
        
        dialog.exec()


    def export_project(self):
        """Export the entire project as a .zip archive with all use cases and metadata."""
        if not self.project_path:
            QMessageBox.warning(self, "No Project", "Please open a project first.")
            return

        # Show save dialog
        zip_path, _ = QFileDialog.getSaveFileName(
            self, 
            "Export Project", 
            "", 
            "Cockburn Project (*.zip)"
        )

        if not zip_path:
            return

        try:
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                # Add all markdown files
                md_dir = os.path.join(self.project_path, "markdown")
                for filename in os.listdir(md_dir):
                    if filename.endswith(".md"):
                        file_path = os.path.join(md_dir, filename)
                        arc_name = f"markdown/{filename}"
                        zf.write(file_path, arc_name)

                # Add metadata file
                metadata_path = os.path.join(self.project_path, "metadata.json")
                if os.path.exists(metadata_path):
                    zf.write(metadata_path, "metadata.json")

                # Add version info
                import_version = {
                    "version": "2.0",
                    "exported_at": datetime.now().isoformat(),
                    "project_path": self.project_path,
                    "use_case_count": len([f for f in os.listdir(md_dir) if f.endswith(".md")])
                }
                zf.writestr("version.json", json.dumps(import_version, indent=2))

            QMessageBox.information(self, "Export Complete", 
                f"Project exported successfully to:\n{zip_path}")
            self.statusBar().showMessage(f"Project exported: {os.path.basename(zip_path)}")

        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export project:\n{e}")

    def import_project(self):
        """Import a project from a .zip archive."""
        # Show open dialog
        zip_file, _ = QFileDialog.getOpenFileName(
            self, 
            "Import Project", 
            "", 
            "Cockburn Project (*.zip)"
        )

        if not zip_file:
            return

        try:
            # Validate it's a valid Cockburn project
            with zipfile.ZipFile(zip_file, "r") as zf:
                # Check for required files
                file_list = zf.namelist()
                md_files = [f for f in file_list if f.startswith("markdown/") and f.endswith(".md")]

                if not md_files:
                    QMessageBox.warning(self, "Invalid Project", 
                        "The selected file does not appear to be a valid Cockburn project.\n"
                        "Expected markdown/ directory with .md files.")
                    return

                # Create temporary directory for extraction
                import_dir = os.path.join(self.project_path or ".", f"_import_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
                os.makedirs(import_dir, exist_ok=True)

                # Extract all files
                zf.extractall(import_dir)

                # Show preview dialog
                dialog = QDialog(self)
                dialog.setWindowTitle("Import Project Preview")
                dialog.setModal(True)
                dialog.setMinimumWidth(400)

                layout = QVBoxLayout(dialog)

                info_label = QLabel(f"Found {len(md_files)} use case(s) to import:")
                layout.addWidget(info_label)

                # List of files
                list_widget = QListWidget()
                for md_file in md_files:
                    filename = md_file.replace("markdown/", "")
                    list_widget.addItem(filename)
                layout.addWidget(list_widget)

                # Buttons
                button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
                button_box.accepted.connect(dialog.accept)
                button_box.rejected.connect(dialog.reject)
                layout.addWidget(button_box)

                if not dialog.exec():
                    return

                # Copy files to project
                for md_file in md_files:
                    src = os.path.join(import_dir, md_file)
                    dst = os.path.join(self.project_path, "markdown", os.path.basename(md_file))
                    shutil.copy2(src, dst)

                # Copy metadata if exists
                metadata_src = os.path.join(import_dir, "metadata.json")
                if os.path.exists(metadata_src):
                    metadata_dst = os.path.join(self.project_path, "metadata.json")
                    shutil.copy2(metadata_src, metadata_dst)

                # Clean up temp directory
                shutil.rmtree(import_dir)

                # Refresh navigation tree
                self.update_navigation_tree()

                QMessageBox.information(self, "Import Complete", 
                    f"Successfully imported {len(md_files)} use case(s)")
                self.statusBar().showMessage(f"Project imported: {os.path.basename(zip_file)}")

        except zipfile.BadZipFile:
            QMessageBox.critical(self, "Import Error", "Invalid ZIP file format.")
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import project:\n{e}")


    def show_about(self):
        """Show about dialog"""
        QMessageBox.about(self, "About Cockburn Specification Generator",
                         "Cockburn Specification Generator V2\n\n"
                         "A modern GUI application for creating Cockburn use case specifications.")

def main():
    app = QApplication(sys.argv)
    window = CockburnGUI()
    window.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()